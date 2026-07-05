"""
pipeline/extractors/docx_extractor.py

Extracts text, tables, headers/footers, embedded images, and core-property
metadata from Microsoft Word (.docx) documents.

Fits into the ingestion pipeline as the extractor registered for .docx files.

Dependency:
  python-docx  → required; lazily imported so the module is always importable
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from pipeline.extractors.text_extractor import (
    BaseExtractor,
    ExtractionError,
    ExtractionResult,
)
from app.models import DocumentMetadata

logger = logging.getLogger(__name__)

# Mapping from Word paragraph style name prefixes to markdown heading markers
_HEADING_STYLES: Dict[str, str] = {
    "heading 1": "#",
    "heading 2": "##",
    "heading 3": "###",
    "heading 4": "####",
    "heading 5": "#####",
    "heading 6": "######",
    "title": "#",
    "subtitle": "##",
}


class DocxExtractor(BaseExtractor):
    """
    Extracts content from .docx files via ``python-docx``.

    What is extracted:
    * Paragraphs — headings are prefixed with Markdown ``#`` based on style.
    * Tables — each cell on one row, rows separated by newlines.
    * Headers and footers (all sections).
    * Embedded image raw bytes.
    * Core properties: author, title, created date, description, keywords.
    """

    supported_extensions = frozenset({".docx"})

    def __init__(self, max_file_size_bytes: Optional[int] = None) -> None:
        from app.config import get_settings
        settings = get_settings()
        self.max_file_size_bytes = max_file_size_bytes or getattr(
            settings, "max_ingest_file_size_bytes", 150 * 1024 * 1024
        )

    # ------------------------------------------------------------------
    # Public interface
    # ------------------------------------------------------------------

    def extract(self, file_path: Path) -> ExtractionResult:
        self._validate_file(file_path, self.max_file_size_bytes)

        try:
            import docx  # noqa: PLC0415
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for .docx extraction. "
                "Install it with: pip install python-docx"
            ) from exc

        try:
            doc = docx.Document(str(file_path))
        except Exception as exc:
            raise ExtractionError(
                f"Could not open DOCX file {file_path}: {exc}"
            ) from exc

        warnings: List[str] = []

        # --- Metadata ---
        metadata = self._build_metadata(file_path, doc)

        # --- Paragraphs (main body) ---
        body_parts: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower()
            prefix = self._heading_prefix(style_name)
            body_parts.append(f"{prefix} {text}" if prefix else text)

        # --- Tables ---
        tables: List[List[List[str]]] = []
        table_text_parts: List[str] = []
        for table in doc.tables:
            rows: List[List[str]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                tables.append(rows)
                table_text_parts.append(self._render_table_text(rows))

        # --- Headers and footers ---
        hf_parts: List[str] = []
        for section in doc.sections:
            for hf in (
                section.header,
                section.footer,
                section.even_page_header,
                section.even_page_footer,
                section.first_page_header,
                section.first_page_footer,
            ):
                if hf and hf.is_linked_to_previous is False:
                    hf_text = "\n".join(
                        p.text.strip() for p in hf.paragraphs if p.text.strip()
                    )
                    if hf_text:
                        hf_parts.append(hf_text)

        # --- Embedded images ---
        images: List[bytes] = []
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        images.append(rel.target_part.blob)
                    except Exception as exc:  # noqa: BLE001
                        warnings.append(f"Skipped image blob: {exc}")
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"Image extraction failed: {exc}")

        all_parts = body_parts + table_text_parts
        if hf_parts:
            all_parts = hf_parts[:1] + all_parts  # Put header text at top

        full_text = "\n\n".join(p for p in all_parts if p)
        if not full_text.strip():
            raise ExtractionError(f"No text content found in {file_path}")

        return ExtractionResult(
            text=full_text,
            metadata=metadata,
            tables=tables,
            images=images,
            warnings=warnings,
        )

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _heading_prefix(self, style_name: str) -> Optional[str]:
        for key, marker in _HEADING_STYLES.items():
            if style_name.startswith(key):
                return marker
        return None

    def _render_table_text(self, rows: List[List[str]]) -> str:
        if not rows:
            return ""
        lines = [" | ".join(rows[0])]
        if len(rows) > 1:
            lines.append(" | ".join("---" for _ in rows[0]))
            lines.extend(" | ".join(row) for row in rows[1:])
        return "\n".join(lines)

    def _build_metadata(self, file_path: Path, doc: Any) -> DocumentMetadata:
        base = self._base_metadata(file_path, "docx")
        extra: Dict[str, Any] = dict(base.extra)

        core = None
        try:
            core = doc.core_properties
        except Exception:  # noqa: BLE001
            pass

        title = (getattr(core, "title", None) or "").strip() or base.title
        author = (getattr(core, "author", None) or "").strip() or None
        description = (getattr(core, "description", None) or "").strip()
        keywords = (getattr(core, "keywords", None) or "").strip()
        revision = getattr(core, "revision", None)

        if description:
            extra["description"] = description
        if keywords:
            extra["keywords"] = keywords
        if revision is not None:
            extra["revision"] = revision

        created_at = getattr(core, "created", None) or base.created_at

        return DocumentMetadata(
            document_id=base.document_id,
            source_path=base.source_path,
            source_type="docx",
            title=title,
            author=author,
            created_at=created_at,
            extra=extra,
        )


__all__ = ["DocxExtractor"]
